from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SCREEN = ROOT / "screenshots"
TODAY = "23 August 2026"
BLUE = RGBColor(31, 78, 121)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 100, 110)
GOLD = RGBColor(171, 117, 21)
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "FFF6DD"
LIGHT_RED = "FDECEC"

def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def rtl_paragraph(p, rtl):
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    if rtl:
        pPr = p._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        pPr.append(bidi)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def cell_margin(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def table(doc, headers, rows, rtl=False, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    for i, header in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, LIGHT_BLUE)
        cell_margin(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        rtl_paragraph(p, rtl)
        r = p.add_run(header)
        set_font(r, size=9.5, color=NAVY, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_margin(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            rtl_paragraph(p, rtl)
            r = p.add_run(str(value))
            set_font(r, size=9.2)
    doc.add_paragraph()
    return t

def note(doc, title, text, rtl=False, tone="blue"):
    fill = LIGHT_GOLD if tone == "gold" else LIGHT_RED if tone == "red" else "F4F6F9"
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, fill)
    cell_margin(c, top=120, start=160, bottom=120, end=160)
    p = c.paragraphs[0]
    rtl_paragraph(p, rtl)
    r = p.add_run(title + "\n")
    set_font(r, size=10.5, color=GOLD if tone == "gold" else NAVY, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5)
    doc.add_paragraph()

def heading(doc, text, level, rtl=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    rtl_paragraph(p, rtl)
    r = p.add_run(text)
    set_font(r, size={1:16,2:13,3:12}[level], color=BLUE if level < 3 else NAVY, bold=True)
    return p

def para(doc, text, rtl=False, bold_prefix=None):
    p = doc.add_paragraph()
    rtl_paragraph(p, rtl)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p

def bullets(doc, items, rtl=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        rtl_paragraph(p, rtl)
        r = p.add_run(item)
        set_font(r)

def steps(doc, items, rtl=False):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        rtl_paragraph(p, rtl)
        r = p.add_run(item)
        set_font(r)

def image(doc, language, filename, caption, rtl=False):
    path = SCREEN / language / "annotated" / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.25))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_font(r, size=9, color=MUTED, italic=True)
    legend = "① Navigation  ② Company/Branch and language context  ③ Current workspace" if not rtl else "① التنقل  ② سياق الشركة والفرع واللغة  ③ مساحة العمل الحالية"
    leg = doc.add_paragraph()
    leg.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.CENTER
    r = leg.add_run(legend)
    set_font(r, size=8.5, color=MUTED)

def add_header_footer(doc, language, rtl=False):
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        r = hp.add_run("DARFUS ERP  |  " + ("دليل المستخدم" if rtl else "End-User Manual"))
        set_font(r, size=9, color=MUTED, bold=True)
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run(("دليل تشغيل المستخدم" if rtl else "Customer operating guide") + "  |  ")
        set_font(r, size=8, color=MUTED)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        fp._p.append(fld)

def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,NAVY,10,5)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

def cover(doc, language, rtl=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("DARFUS ERP")
    set_font(r, size=28, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("دليل المستخدم" if rtl else "End-User Manual")
    set_font(r, size=22, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("دليل التشغيل اليومي للمخزون والمبيعات والمالية" if rtl else "A practical guide to inventory, sales, and finance")
    set_font(r, size=13, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    r = p.add_run(("اللغة: العربية" if rtl else "Language: English") + "\nVersion: Stage F runtime handover\n" + TODAY)
    set_font(r, size=11, color=MUTED)
    note(doc, "مهم" if rtl else "Important", "هذا الدليل يشرح ما يظهر في الإصدار الحالي. قد تختلف الأزرار حسب الدور والفرع وحالة السجل." if rtl else "This guide describes the current release. Buttons may differ by role, branch, and record status.", rtl, "gold")
    doc.add_page_break()

def toc(doc, rtl=False):
    heading(doc, "المحتويات" if rtl else "Contents", 1, rtl)
    chapters = ([
        "أين أبدأ؟", "أساسيات التطبيق", "لوحة التحكم", "المخزون والاستلام", "ملفات المخزون",
        "الأصل والباركود وRFID", "التحويل والورشة والجرد", "الموردون والمشتريات", "العملاء", "نقطة البيع",
        "الفواتير والمرتجعات والحجوزات", "مركز الذهب", "شراء الذهب من العميل", "الحسابات والخزنة والتقارير",
        "المستخدمون والإعدادات", "الأخطاء والسلامة", "قاموس الحالات والمهام", "قائمة التسليم والتدريب",
    ] if rtl else [
        "Where do I start?", "Application basics", "Dashboard", "Inventory and receiving", "Inventory profiles",
        "Asset, Barcode and RFID", "Transfers, Workshop and Count", "Suppliers and purchases", "Customers", "Point of Sale",
        "Invoices, returns and reservations", "Gold Center", "Customer Gold Purchase", "Accounting, Treasury and Reports",
        "Users and Settings", "Troubleshooting and safety", "Status glossary and task finder", "Handover and training checklist",
    ])
    for i, chapter in enumerate(chapters, 1):
        para(doc, f"{i}. {chapter}", rtl)
    note(doc, "طريقة القراءة" if rtl else "How to use this guide", "ابدأ بفصل الدور المناسب لك، ثم استخدم وصف المهمة السريعة قبل الدخول في التفاصيل." if rtl else "Start with the persona that matches your work, then use the quick task recipes before reading the detailed chapters.", rtl)
    doc.add_page_break()

def persona_section(doc, rtl=False):
    heading(doc, "أين أبدأ؟" if rtl else "Where do I start?", 1, rtl)
    para(doc, "اختر الوصف الأقرب إلى عملك. قد تظهر لك شاشة أو أزرار أقل من الأمثلة إذا لم يمنحك دورك صلاحية العملية." if rtl else "Choose the description closest to your work. You may see fewer screens or buttons than the examples when your role does not include the required permission.", rtl)
    rows = [
        (("الكاشير", "POS، العملاء، الفواتير"), ("Cashier", "POS, Customers, Invoices"), "البيع، البحث بالباركود، مراجعة الإجمالي وطباعة الإيصال." if rtl else "Sell, search by Barcode, review totals, and print the receipt."),
        (("موظف المخزون", "المخزون، الاستلام، التحويل، الورشة، الجرد"), ("Inventory / Warehouse", "Inventory, Receive, Transfer, Workshop, Count"), "تتبع القطع، المواقع، الحالة والباركود." if rtl else "Track pieces, locations, status, and Barcode."),
        (("المشتريات", "الموردون، الاستلام، الرصيد"), ("Purchasing", "Suppliers, Receive, Supplier balance"), "استلام القطع من المورد ومتابعة سجل الشراء." if rtl else "Receive supplier stock and follow purchase history."),
        (("المحاسب", "الحسابات، القوائم، الخزنة"), ("Accountant", "Accounting, Statements, Treasury"), "مراجعة القيود، الأرصدة والخزنة." if rtl else "Review journals, balances, and treasury."),
        (("مدير الفرع", "لوحة التحكم، التقارير، الموافقات"), ("Branch Manager", "Dashboard, Reports, Approvals"), "متابعة الأداء والموافقات وحالة الفرع." if rtl else "Monitor performance, approvals, and branch status."),
        (("مدير النظام", "المستخدمون، الصلاحيات، الإعدادات"), ("Administrator", "Users, Permissions, Settings"), "إدارة الوصول والبيانات المرجعية والإعدادات الحساسة." if rtl else "Manage access, master data, and sensitive settings."),
    ]
    headers = ["الدور / المجال", "أهم الشاشات", "الاستخدام اليومي"] if rtl else ["Persona", "Main screens", "Daily use"]
    data=[]
    for ar,en,use in rows:
        data.append((ar[0] if rtl else en[0], ar[1] if rtl else en[1], use))
    table(doc, headers, data, rtl, [1.4,2.1,3.0])
    note(doc, "صلاحية" if rtl else "Permission note", "عدم ظهور الزر لا يعني أن الشاشة معطلة؛ تحقق من الدور والفرع وحالة السجل." if rtl else "A missing button does not necessarily mean the screen is broken; check your role, branch, and record status.", rtl, "gold")

def basics_section(doc, rtl=False):
    heading(doc, "أساسيات التطبيق" if rtl else "Application basics", 1, rtl)
    image(doc, "ar" if rtl else "en", "AR-01-dashboard.png" if rtl else "EN-01-dashboard.png", "لوحة التحكم وسياق الفرع" if rtl else "Dashboard and branch context", rtl)
    heading(doc, "تسجيل الدخول واللغة والفرع" if rtl else "Login, language, and branch", 2, rtl)
    steps(doc, [
        "افتح صفحة الدخول وأدخل بيانات الحساب التي سلّمها لك مدير النظام." if rtl else "Open the login page and enter the account details provided by your administrator.",
        "بعد الدخول، راجع اسم الشركة والفرع الظاهرين أعلى الشاشة قبل أي عملية." if rtl else "After signing in, check the company and branch shown in the header before starting any action.",
        "استخدم زر اللغة للتبديل بين العربية والإنجليزية. اتجاه الصفحة يتغير تلقائيًا." if rtl else "Use the language control to switch between Arabic and English. The page direction changes automatically.",
        "استخدم زر الفرع لاختيار الفرع المسموح لحسابك؛ تغيير الفرع لا ينقل قطعة تلقائيًا." if rtl else "Use the branch control to choose a branch available to your account; changing the branch does not move an Asset.",
    ], rtl)
    heading(doc, "ما الذي تراه في معظم الشاشات؟" if rtl else "What you see on most screens", 2, rtl)
    bullets(doc, [
        "الشريط الجانبي: الانتقال بين الوحدات." if rtl else "Sidebar: move between modules.",
        "الشريط العلوي: البحث العام، الشركة، الفرع، اللغة، التنبيهات وحساب المستخدم." if rtl else "Top bar: global search, company, branch, language, notifications, and user account.",
        "العناوين والأشرطة: توضح الحالة أو عدد النتائج أو صلاحية الإجراء." if rtl else "Headings and badges: show the screen, result count, status, or action availability.",
        "الجداول: استخدم البحث والمرشحات والصفحات، ولا تعتبر الصفحة الحالية كل البيانات." if rtl else "Tables: use search, filters, and pagination; the current page may not contain every record.",
    ], rtl)

def inventory_section(doc, language, rtl=False):
    heading(doc, "المخزون والاستلام" if rtl else "Inventory and receiving", 1, rtl)
    image(doc, language, "AR-04-inventory-intake-chooser.png" if rtl else "EN-04-inventory-intake-chooser.png", "اختيار نوع المخزون من المسار الموحد" if rtl else "Choosing a profile from the unified intake path", rtl)
    para(doc, "المسار المعتمد هو: المخزون ← إضافة / استلام مخزون ← اختر النوع ← أكمل النموذج ← راجع المعاينة ← استلام المخزون." if rtl else "The canonical path is: Inventory -> Add / Receive Inventory -> choose the profile -> complete the form -> review the preview -> Receive Inventory.", rtl)
    heading(doc, "ما الذي يعنيه الاستلام؟" if rtl else "What receiving means", 2, rtl)
    bullets(doc, [
        "كل قطعة فعلية تسجل كأصل مستقل، ويحافظ الأصل على هويته وباركوده." if rtl else "Each physical piece is recorded as its own Asset, with its identity and Barcode preserved.",
        "يظهر الأصل في قائمة المخزون وفي تفاصيل الأصل حسب الشركة والفرع والموقع." if rtl else "The Asset appears in Inventory and Asset Details according to company, branch, and location.",
        "يسجل النظام بيانات الشراء والمورد والتاريخ والقيمة التاريخية عند الحاجة." if rtl else "Purchase, supplier, date, and historical value are recorded when the workflow requires them.",
        "قد ينتج عن الاستلام أثر في رصيد المورد والقيد المالي حسب نوع العملية وإعداداتها." if rtl else "The receive may affect supplier balance and accounting according to the transaction type and settings.",
        "لا تستخدم شاشة الموردين القديمة لإنشاء استلام إذا كانت شاشة المخزون الموحدة متاحة." if rtl else "Do not use an old supplier shortcut to create a receive when the unified Inventory path is available.",
    ], rtl)
    heading(doc, "خطوات الاستلام" if rtl else "Receiving steps", 2, rtl)
    steps(doc, [
        "تأكد من الفرع والموقع والمورد وتاريخ الشراء." if rtl else "Confirm the branch, location, supplier, and purchase date.",
        "اختر ملف المخزون المناسب ولا تخلط بين القطعة والوزن أو الحجر الحر والمجوهرات." if rtl else "Choose the correct inventory profile; do not mix piece, weight, loose-stone, and jewellery meanings.",
        "أدخل الحقول المطلوبة فقط بالقيم الفعلية، ثم راجع المعاينة والأسعار والضريبة." if rtl else "Enter the required fields with the actual values, then review the preview, prices, and tax.",
        "اضغط استلام المخزون مرة واحدة بعد المراجعة." if rtl else "Press Receive Inventory once after reviewing the form.",
    ], rtl)
    note(doc, "بعد التنفيذ" if rtl else "What happens after completion", "يظهر نجاح العملية، وتستطيع العثور على الأصل في المخزون وتفاصيل الأصل وسجل المورد والقيد عند توفره. لا تكرر الضغط إذا كانت النتيجة غير واضحة؛ تواصل مع المسؤول." if rtl else "A success result appears, and you can find the Asset in Inventory, Asset Details, supplier history, and accounting when applicable. Do not repeat the action when the result is uncertain; contact the administrator.", rtl, "gold")

def profile_section(doc, language, rtl=False):
    heading(doc, "ملفات المخزون" if rtl else "Inventory profiles", 1, rtl)
    profiles = [
        ("ذهب بالوزن", "Gold By Weight", "AR-02-inventory-gold-by-weight.png" if rtl else "EN-02-inventory-gold-by-weight.png", "استخدمه عندما تكون قيمة الذهب الأساسية مرتبطة بالوزن والعيار." if rtl else "Use it when the gold value is driven by weight and karat."),
        ("ذهب بالقطعة", "Gold By Piece", "AR-02-inventory-gold-by-piece.png" if rtl else "EN-02-inventory-gold-by-piece.png", "استخدمه عندما تمثل كل قطعة وحدة ذهب منفصلة ببيانات القطعة." if rtl else "Use it when each gold piece is a separate physical unit with piece-specific data."),
        ("مجوهرات ألماس", "Diamond Jewellery", "AR-02-inventory-diamond-jewellery.png" if rtl else "EN-02-inventory-diamond-jewellery.png", "استخدمه للمجوهرات المركبة التي تحتوي على ألماس." if rtl else "Use it for jewellery that contains mounted diamonds."),
        ("ألماس حر", "Loose Diamond", "AR-02-inventory-loose-diamond.png" if rtl else "EN-02-inventory-loose-diamond.png", "كل حجر حر يسجل كقطعة مستقلة بباركود خاص به." if rtl else "Each loose stone is recorded as an individual piece with its own Barcode."),
        ("مجوهرات الأحجار الكريمة", "Gem Stone Jewellery", "AR-02-inventory-gem-stone.png" if rtl else "EN-02-inventory-gem-stone.png", "استخدمه للمجوهرات التي تحتوي على أحجار كريمة مركبة." if rtl else "Use it for jewellery with mounted gemstones."),
        ("حجر كريم حر", "Loose Gem Stone", "AR-02-inventory-loose-gem-stone.png" if rtl else "EN-02-inventory-loose-gem-stone.png", "استخدمه للحجر الحر الذي يتم تتبعه منفردًا." if rtl else "Use it for a loose gemstone tracked separately."),
        ("مجوهرات اللؤلؤ", "Pearl Jewellery", "AR-02-inventory-pearl.png" if rtl else "EN-02-inventory-pearl.png", "استخدمه للمجوهرات التي تحتوي على لؤلؤ مركب." if rtl else "Use it for jewellery containing mounted pearls."),
        ("لؤلؤ منفرد", "Loose Pearl", "AR-02-inventory-loose-pearl.png" if rtl else "EN-02-inventory-loose-pearl.png", "استخدمه للؤلؤة المنفردة التي تحتاج أصلًا وباركودًا مستقلين." if rtl else "Use it for an individual pearl that needs its own Asset and Barcode."),
    ]
    for ar_name, en_name, shot, meaning in profiles:
        heading(doc, ar_name if rtl else en_name, 2, rtl)
        image(doc, language, shot, f"{ar_name}" if rtl else f"{en_name}", rtl)
        para(doc, meaning, rtl)
        bullets(doc, [
            "راجع بيانات المورد والموقع والتاريخ والضريبة المشتركة في أعلى النموذج." if rtl else "Review the shared supplier, location, date, and tax information at the top of the form.",
            "أدخل بيانات الملف التي تظهر على الشاشة؛ لا تضع قيمة في حقل غير معروض أو غير مطلوب." if rtl else "Enter the profile fields shown on the screen; do not invent a value for a field that is not shown or required.",
            "راجع المعاينة التاريخية والحالية وسعر البيع قبل فتح التأكيد." if rtl else "Review the historical and current preview and selling price before opening confirmation.",
            "بعد الاستلام، ابحث عن الأصل من المخزون أو الباركود." if rtl else "After receiving, find the Asset from Inventory or by Barcode.",
        ], rtl)
        note(doc, "أثر الملف" if rtl else "Profile effect", "يظل الأصل والباركود وسجل الشراء مرتبطين بالقطعة نفسها. القيم المحسوبة أو الضريبة تعرضها الشاشة بحسب إعدادات الشركة." if rtl else "The Asset, Barcode, and purchase history remain linked to the same physical piece. Calculated values and tax are shown according to company settings.", rtl)

def asset_section(doc, language, rtl=False):
    heading(doc, "الأصل والباركود وRFID" if rtl else "Asset, Barcode, and RFID", 1, rtl)
    image(doc, language, "AR-05-asset-details.png" if rtl else "EN-05-asset-details.png", "تفاصيل الأصل وسجل التتبع" if rtl else "Asset details and traceability", rtl)
    para(doc, "الأصل هو بطاقة التتبع التجارية للقطعة الفعلية. استخدم تفاصيل الأصل لمراجعة الهوية والفرع والموقع والحالة والمورد والتاريخ والقيمة وسجل الأحداث." if rtl else "Asset Details is the business traceability card for a physical piece. Use it to review identity, branch, location, status, supplier, history, values, and events.", rtl)
    heading(doc, "الباركود" if rtl else "Barcode", 2, rtl)
    bullets(doc, [
        "البحث بالباركود يعيد القطعة ذات الهوية المطابقة فقط." if rtl else "Barcode search returns the matching physical item.",
        "إعادة الطباعة تعيد نفس الهوية ولا تنشئ قطعة جديدة." if rtl else "Reprint uses the same identity; it does not create a new piece.",
        "استبدال الباركود - إذا ظهر الزر وكان مسموحًا - يغير الباركود النشط للقطعة نفسها ويترك السجل السابق للتتبع." if rtl else "Barcode replacement - when exposed and permitted - changes the active Barcode for the same Asset and keeps the previous history for traceability.",
    ], rtl)
    heading(doc, "RFID" if rtl else "RFID", 2, rtl)
    para(doc, "وثّق استخدام RFID فقط إذا ظهر زر أو حقل RFID في حسابك. ربط الوسم أو تغييره لا يغير هوية الأصل أو الباركود إلا إذا أوضحت الشاشة ذلك صراحة." if rtl else "Use RFID only when the RFID control is visible to your account. Assigning or replacing a tag does not change the Asset or Barcode unless the screen explicitly says so.", rtl)
    note(doc, "لا تعدل الحالة يدويًا" if rtl else "Do not edit status by imitation", "تغيير الحالة يجب أن يتم من خلال المسار الصحيح مثل البيع أو التحويل أو الورشة، وليس بتعديل عشوائي في تفاصيل الأصل." if rtl else "Status changes belong to the correct workflow such as Sale, Transfer, or Workshop, not arbitrary editing in Asset Details.", rtl, "gold")

def lifecycle_section(doc, language, rtl=False):
    heading(doc, "التحويل والورشة والجرد" if rtl else "Transfers, Workshop, and Count", 1, rtl)
    heading(doc, "تحويل الفروع" if rtl else "Branch Transfers", 2, rtl)
    image(doc, language, "AR-01-inventory-transfers.png" if rtl else "EN-01-inventory-transfers.png", "سجل التحويلات ومراحلها" if rtl else "Transfer register and states", rtl)
    steps(doc, ["المخزون ← تحويلات الفروع ← إنشاء تحويل." if rtl else "Inventory -> Transfers -> Create Transfer.", "حدد فرع وموقع المصدر والوجهة، ثم اختر الأصول المتاحة." if rtl else "Choose source and destination branch/location, then select available Assets.", "أرسل الطلب ثم نفذ الموافقة والإرسال والاستلام حسب الدور." if rtl else "Submit the request, then complete Approve, Dispatch, and Receive according to the role.", "افتح السجل نفسه لمراجعة الحالة." if rtl else "Open the same transfer record to review its state."], rtl)
    note(doc, "ماذا يتغير؟" if rtl else "What changes?", "تبقى القطعة والأصل والباركود كما هي، ويتغير الفرع أو الموقع عبر عملية التحويل. لا ينشئ التحويل فاتورة بيع أو دفعة مورد." if rtl else "The physical piece, Asset, and Barcode remain the same; branch/location custody changes through the transfer. A transfer does not create a sales invoice or supplier payment.", rtl)
    heading(doc, "الورشة" if rtl else "Workshop", 2, rtl)
    image(doc, language, "AR-01-inventory-workshop.png" if rtl else "EN-01-inventory-workshop.png", "سجل الورشة" if rtl else "Workshop register", rtl)
    steps(doc, ["المخزون ← الورشة ← اختر القطعة والوجهة." if rtl else "Inventory -> Workshop -> choose the piece and destination.", "أكد الإرسال، ثم تابع السجل نفسه عند الإرجاع." if rtl else "Confirm Send, then follow the same record when it is returned.", "لا تبيع القطعة أثناء وجودها في الورشة إذا كانت حالتها تمنع البيع." if rtl else "Do not sell the piece while it is in Workshop when its status prevents sale."], rtl)
    note(doc, "ماذا يحدث؟" if rtl else "What happens?", "يبقى الأصل والباركود كما هما، وتتغير الحيازة التشغيلية والحالة حسب مسار الورشة. لا تعتبر الورشة بيعًا." if rtl else "The Asset and Barcode stay the same while operational custody/status changes through Workshop. Workshop is not a sale.", rtl)
    heading(doc, "الجرد" if rtl else "Inventory Count", 2, rtl)
    image(doc, language, "AR-01-inventory-stock-audit.png" if rtl else "EN-01-inventory-stock-audit.png", "الجرد ومقارنة المتوقع بالمعدود" if rtl else "Count and compare expected with counted", rtl)
    steps(doc, ["المخزون ← جرد المخزون ← اختر الموقع وابدأ جردًا." if rtl else "Inventory -> Inventory Count -> choose the location and start a count.", "راجع Expected وCounted وMissing وUnexpected وVariance." if rtl else "Review Expected, Counted, Missing, Unexpected, and Variance.", "أكمل ثم أغلق نفس سجل الجرد إذا كان المسار يتطلب ذلك." if rtl else "Complete and close the same count when the workflow requires it."], rtl)
    note(doc, "أثر الجرد" if rtl else "Count effect", "الجرد يسجل دليل ما تم رصده. لا تفترض أنه يغير حالة الأصل تلقائيًا؛ راجع نتيجة الجرد وسجل التدقيق." if rtl else "A count records evidence of what was observed. Do not assume it changes Asset status automatically; review the count result and audit history.", rtl)

def suppliers_section(doc, language, rtl=False):
    heading(doc, "الموردون والمشتريات" if rtl else "Suppliers and purchases", 1, rtl)
    image(doc, language, "AR-05-supplier-details.png" if rtl else "EN-05-supplier-details.png", "بيانات المورد ورصيده وسجله" if rtl else "Supplier details, balance, and history", rtl)
    steps(doc, ["افتح الموردين والمشتريات وابحث بالاسم أو الفئة أو الهاتف." if rtl else "Open Suppliers & purchases and search by name, category, or phone.", "افتح عرض المورد لمراجعة بياناته وسجل المشتريات والرصيد." if rtl else "Open View to review supplier details, purchase history, and balance.", "لإنشاء أو تعديل مورد استخدم الزر الظاهر لحسابك فقط وأكمل الحقول المعروضة." if rtl else "For create or edit, use the button visible to your account and complete the fields shown.", "لا تكتب الرصيد يدويًا؛ الرصيد المحسوب يظهر من معاملات النظام." if rtl else "Do not overwrite the balance manually; the system-derived balance comes from recorded transactions."], rtl)
    heading(doc, "استلام مشتريات المورد" if rtl else "Receiving supplier purchases", 2, rtl)
    para(doc, "يبدأ الاستلام من المخزون عبر المسار الموحد. شاشة المورد للمراجعة والبيانات والتاريخ، وليست مصدرًا ثانيًا لإنشاء استلام عندما تكون شاشة المخزون متاحة." if rtl else "Start receiving from Inventory through the unified path. Supplier Details is for information, history, and balance; it is not a second receive source when Inventory is available.", rtl)
    note(doc, "بعد الدفع أو الاستلام" if rtl else "After receiving or paying", "بعد الاستلام راجع الأصل في المخزون، وسجل المورد، والحسابات عند الحاجة. بعد دفعة مورد ناجحة راجع الرصيد والخزنة والقيد." if rtl else "After a receive, review the Asset in Inventory, supplier history, and Accounting when applicable. After a successful supplier payment, review the balance, Treasury, and journal.", rtl)

def customers_section(doc, language, rtl=False):
    heading(doc, "العملاء" if rtl else "Customers", 1, rtl)
    image(doc, language, "AR-07-customer-details.png" if rtl else "EN-07-customer-details.png", "تفاصيل العميل" if rtl else "Customer details", rtl)
    steps(doc, ["العملاء وCRM ← إضافة عميل." if rtl else "Customers & CRM -> Add Customer.", "أدخل الحقول المطلوبة كما تظهر، ثم احفظ." if rtl else "Complete the required fields shown, then save.", "ابحث بالاسم أو الهاتف وافتح التفاصيل عند الحاجة." if rtl else "Search by name or phone and open details when needed.", "في POS ابحث بالهاتف ثم اختر العميل؛ لا ينشئ POS عميلًا مجهولًا تلقائيًا." if rtl else "In POS, search by phone and select the customer; POS does not silently create an unknown customer."], rtl)
    note(doc, "أثر العميل" if rtl else "Customer effect", "إنشاء أو تعديل بيانات العميل لا ينشئ فاتورة أو دفعة. البيع أو العربون قد يظهر في ملخص العميل حسب المسار." if rtl else "Creating or editing a customer does not create an invoice or payment. Sales or deposits may appear in the customer summary according to the workflow.", rtl)

def pos_section(doc, language, rtl=False):
    heading(doc, "نقطة البيع" if rtl else "Point of Sale", 1, rtl)
    image(doc, language, "AR-04-pos-cart-journal-preview.png" if rtl else "EN-04-pos-cart-journal-preview.png", "السلة والإجماليات قبل الإتمام" if rtl else "Cart and totals before completion", rtl)
    heading(doc, "بيع قطعة" if rtl else "Making a sale", 2, rtl)
    steps(doc, ["تأكد من الفرع والعميل الظاهرين." if rtl else "Confirm the branch and customer shown.", "ابحث بالباركود أو معرّف القطعة، ثم اختر الأصل المتاح." if rtl else "Search by Barcode or supported item identifier, then select an available Asset.", "راجع السلة والسعر والتوفر والضريبة والإجمالي." if rtl else "Review the cart, price, availability, VAT, and total.", "راجع معاينة القيد إن ظهرت." if rtl else "Review Journal Preview when it is shown.", "اختر طريقة الدفع المعروضة، ثم افتح التأكيد وأكمل Checkout مرة واحدة." if rtl else "Choose a displayed payment method, open confirmation, and complete Checkout once.", "افتح أو اطبع الإيصال بعد النجاح." if rtl else "Open or print the receipt after success."], rtl)
    heading(doc, "معاينة القيد" if rtl else "Journal Preview", 2, rtl)
    image(doc, language, "AR-04-pos-journal-details.png" if rtl else "EN-04-pos-journal-details.png", "تفاصيل المعاينة المالية" if rtl else "Financial preview details", rtl)
    para(doc, "معاينة القيد تعرض الأثر المالي المتوقع قبل الإتمام. في المثال الظاهر: النقدية والإيراد والضريبة وتكلفة البضاعة والمخزون؛ ظهور القيد متزنًا يعني أن العرض متوازن قبل Checkout، وليس أن البيع تم." if rtl else "Journal Preview shows the expected financial effect before completion. The example shows cash, revenue, VAT, cost of goods sold, and inventory; a balanced preview means the preview balances, not that the sale is complete.", rtl)
    note(doc, "بعد Checkout" if rtl else "What happens after Checkout", "ينشأ مستند البيع والإيصال، تسجل الدفعة، يتأثر النقد أو الحساب المناسب، يرحل القيد، وتتغير حالة الأصل إلى Sold مع بقاء الباركود مرتبطًا بالأصل نفسه." if rtl else "The sale document and receipt are created, payment is recorded, cash or the relevant account is affected, the journal is posted, and the Asset becomes Sold while its Barcode remains linked to the same Asset.", rtl, "gold")
    note(doc, "نتيجة غير واضحة" if rtl else "Uncertain result", "إذا ظهر خطأ أو انتهت العملية دون نتيجة واضحة، لا تعاود Checkout أو الدفع عشوائيًا؛ تواصل مع المشرف ليتحقق من الفاتورة والدفعة." if rtl else "If an error or uncertain result appears, do not repeat Checkout or payment blindly; ask a supervisor to verify the invoice and payment.", rtl, "red")

def sales_section(doc, language, rtl=False):
    heading(doc, "الفواتير والمرتجعات والحجوزات" if rtl else "Invoices, returns, and reservations", 1, rtl)
    image(doc, language, "AR-07-invoice-details.png" if rtl else "EN-07-invoice-details.png", "تفاصيل فاتورة موجودة" if rtl else "Existing invoice details", rtl)
    heading(doc, "الفواتير والطباعة" if rtl else "Invoices and print", 2, rtl)
    steps(doc, ["المبيعات والفواتير ← بحث وطباعة الفواتير." if rtl else "Invoices & sales -> Invoices Search & Print.", "ابحث برقم الفاتورة أو العميل أو طريقة الدفع." if rtl else "Search by invoice, customer, or payment method.", "افتح عرض لمراجعة العميل والقطعة والضريبة والإجمالي وحالة الدفع." if rtl else "Open View to review customer, item, VAT, total, and payment status.", "استخدم طباعة الفاتورة إذا ظهر الزر." if rtl else "Use Print invoice when the button is shown."], rtl)
    para(doc, "بيانات العميل الظاهرة داخل الفاتورة تمثل لقطة وقت البيع، وقد تختلف لاحقًا عن بيانات العميل الحالية." if rtl else "Customer information in an invoice represents the sale-time snapshot and may differ later from the current customer profile.", rtl)
    heading(doc, "المرتجع والاستبدال" if rtl else "Returns and exchanges", 2, rtl)
    para(doc, "ابدأ من شاشة المرتجعات أو الاستبدال، ابحث عن الفاتورة الأصلية، وتابع التحقق والحالة التي تسمح بها الشاشة. لا تستخدم تعديل الحالة اليدوي بدل المسار." if rtl else "Start from Returns or Exchanges, locate the original invoice, and follow the validations and state allowed by the screen. Do not imitate the workflow by editing status manually.", rtl)
    heading(doc, "الحجز والعربون" if rtl else "Reservations and deposits", 2, rtl)
    steps(doc, ["المبيعات ← الحجوزات ← اختر العميل والأصل." if rtl else "Sales -> Reservations -> choose the customer and Asset.", "سجل العربون فقط إذا ظهر الزر وكانت الصلاحية متاحة." if rtl else "Record a deposit only when the button is shown and the permission is available.", "تابع الإيداع أو الاسترداد أو إكمال البيع من نفس سجل الحجز." if rtl else "Follow deposit, refund, or complete sale from the same reservation record."], rtl)
    note(doc, "أثر العربون" if rtl else "Deposit effect", "العربون يظل مقدمًا للعميل حتى الاسترداد أو إكمال البيع، ويؤثر في الخزنة والحساب حسب نوع العملية." if rtl else "A deposit remains a customer advance until refund or completion and affects Treasury and Accounting according to the action.", rtl)

def gold_cgp_section(doc, language, rtl=False):
    heading(doc, "مركز الذهب وشراء الذهب من العميل" if rtl else "Gold Center and Customer Gold Purchase", 1, rtl)
    image(doc, language, "AR-01-gold-center.png" if rtl else "EN-01-gold-center.png", "حالة السعر ومعدل التحديث" if rtl else "Price status and update state", rtl)
    heading(doc, "مركز الذهب" if rtl else "Gold Center", 2, rtl)
    para(doc, "يعرض مركز الذهب الأسعار المرجعية الحالية والعيارات وحالة التحديث. السعر المرجعي يستخدمه المسار الذي يدعمه، ولا ينبغي تجاوز حالة قديمة أو غير متاحة." if rtl else "Gold Center shows current reference prices, karats, and freshness state. Supported pricing workflows use this reference; do not bypass a stale or unavailable state.", rtl)
    heading(doc, "شراء الذهب من العميل" if rtl else "Customer Gold Purchase", 2, rtl)
    image(doc, language, "AR-01-sales-customer-gold-drafts.png" if rtl else "EN-01-sales-customer-gold-drafts.png", "مسودات شراء الذهب من العميل" if rtl else "Customer Gold Purchase drafts", rtl)
    steps(doc, ["افتح شراء الذهب من العميل وأنشئ مسودة إذا ظهر الزر." if rtl else "Open Customer Gold Purchase and create a draft when the button is shown.", "أدخل بيانات العميل والذهب، ثم راجع التحقق والسعر واللقطة." if rtl else "Enter customer and gold details, then review validation, price, and snapshot.", "انتقل إلى الترحيل فقط بعد التحقق، ثم راجع التسوية إذا كانت متاحة." if rtl else "Post only after validation, then review Settlement when available."], rtl)
    note(doc, "ليس استلام مورد" if rtl else "Not a supplier receive", "شراء الذهب من العميل مسار مستقل عن استلام المخزون من المورد. لا تخلط بين المستندين." if rtl else "Customer Gold Purchase is separate from supplier stock receiving. Do not mix the two documents.", rtl, "gold")

def finance_section(doc, language, rtl=False):
    heading(doc, "الحسابات والخزنة والتقارير" if rtl else "Accounting, Treasury, and Reports", 1, rtl)
    image(doc, language, "AR-01-accounting.png" if rtl else "EN-01-accounting.png", "شاشة الحسابات" if rtl else "Accounting screen", rtl)
    bullets(doc, [
        "الحسابات: راجع القيود والمصدر والتاريخ والمدين والدائن والحالة." if rtl else "Accounting: review journals, source, date, debit, credit, and status.",
        "دليل الحسابات: راجع الحسابات وربط الفرع عندما تكون الصلاحية متاحة." if rtl else "Chart of Accounts: review accounts and branch mappings when permitted.",
        "القوائم المالية والتقارير: استخدم البحث والفترة والفرع، ولا تفسر شاشة فارغة كأنها صفر مالي نهائي قبل التأكد من المرشحات." if rtl else "Statements and Reports: use date and branch filters; do not treat an empty screen as a final zero before checking filters.",
        "الخزنة: راجع جلسة الخزنة والحركات ونتيجة المدفوعات حسب ما يظهر لحسابك." if rtl else "Treasury: review cash session, movements, and payment results shown to your account.",
    ], rtl)
    image(doc, language, "AR-01-accounting-treasury.png" if rtl else "EN-01-accounting-treasury.png", "الخزنة والحركات" if rtl else "Treasury and movements", rtl)
    note(doc, "حماية القيود" if rtl else "Journal safety", "لا تعدل قيدًا مرحلًا يدويًا. إذا كان القيد أو الدفع غير واضح، احتفظ بالمرجع واطلب مراجعة المسؤول المالي." if rtl else "Do not manually rewrite a posted journal. If a journal or payment is uncertain, keep the reference and ask the finance administrator to review it.", rtl, "gold")

def settings_section(doc, language, rtl=False):
    heading(doc, "المستخدمون والإعدادات" if rtl else "Users and Settings", 1, rtl)
    image(doc, language, "AR-01-settings.png" if rtl else "EN-01-settings.png", "إعدادات الشركة" if rtl else "Company settings", rtl)
    bullets(doc, [
        "المستخدمون والموظفون: إنشاء الحساب أو تعطيله أو ضبط الدور والفرع إذا ظهرت الصلاحية." if rtl else "Users and Employees: create, deactivate, and assign role/branch when permitted.",
        "الضريبة وضريبة القيمة المضافة: راجع سياسة الشركة ومعدلها ولا تغيرها إلا من المسؤول المختص." if rtl else "Tax & VAT: review company policy and rate; only the authorized administrator should change them.",
        "المواقع: أضف أو عطّل الموقع من شاشة المواقع عندما يكون ذلك مسموحًا، ولا تحذف قيمة مستخدمة تاريخيًا." if rtl else "Locations: add or disable a location when permitted; do not delete a value used by historical records.",
        "أكواد الباركود: راجع الأكواد من الشاشة المخصصة ولا تنشئ كودًا يدويًا خارج المسار." if rtl else "Barcode codes: review codes in the dedicated screen; do not create codes outside the supported workflow.",
    ], rtl)
    image(doc, language, "AR-02-settings-tax.png" if rtl else "EN-02-settings-tax.png", "سياسة الضريبة" if rtl else "Tax policy", rtl)
    note(doc, "إعدادات المال" if rtl else "Money settings", "الإعدادات التي تؤثر في الأسعار أو الضريبة أو الحسابات يجب أن يغيرها المسؤول المخول فقط، مع توثيق سبب التغيير." if rtl else "Settings that affect price, tax, or accounting should be changed only by the authorized administrator, with a documented reason.", rtl, "red")

def safety_section(doc, rtl=False):
    heading(doc, "الأخطاء والسلامة" if rtl else "Troubleshooting and safety", 1, rtl)
    rows = [
        ("الزر غير ظاهر", "راجع الدور والفرع وحالة السجل." if rtl else "Check role, branch, and record status."),
        ("Permission denied", "اطلب من المسؤول مراجعة الصلاحية." if rtl else "Ask the administrator to review the permission."),
        ("الباركود غير موجود", "تأكد من الكتابة والفرع وحالة القطعة." if rtl else "Check spelling, branch, and item status."),
        ("العنصر غير متاح للبيع", "قد يكون مباعًا أو محجوزًا أو في ورشة أو تحويل." if rtl else "It may be Sold, Reserved, in Workshop, or in Transfer."),
        ("السعر غير متاح أو قديم", "لا تتجاوز حماية السعر؛ راجع مركز الذهب أو المسؤول." if rtl else "Do not bypass price protection; check Gold Center or contact the administrator."),
        ("المعاينة غير متاحة", "أكمل الحقول المطلوبة وتحقق من الفرع والموقع والإعدادات." if rtl else "Complete required fields and check branch, location, and settings."),
        ("نتيجة دفع أو Checkout غير واضحة", "لا تعاود العملية؛ اطلب فحص الفاتورة والدفعة." if rtl else "Do not repeat it; ask for invoice and payment verification."),
        ("الجلسة انتهت", "سجل الدخول من جديد ولا تعيد الإجراء المالي قبل التحقق." if rtl else "Sign in again and verify any financial action before repeating it."),
    ]
    table(doc, ["ما تراه" if rtl else "What you see", "ما الذي تفعله" if rtl else "What to do"], rows, rtl, [2.0,4.5])
    heading(doc, "لا تفعل ذلك" if rtl else "Do not do this", 2, rtl)
    bullets(doc, [
        "لا تبع القطعة الفعلية نفسها مرتين." if rtl else "Do not sell the same physical piece twice.",
        "لا تنقل القطعة بين الفروع خارج التحويل." if rtl else "Do not move an item between branches outside Transfer.",
        "لا تنشئ موردًا أو عميلًا وهميًا لتجاوز حقل مطلوب." if rtl else "Do not create a fake supplier or customer to bypass a required field.",
        "لا تعاود الدفع أو Checkout دون معرفة نتيجة المحاولة الأولى." if rtl else "Do not repeat payment or Checkout without knowing the result of the first attempt.",
        "لا تحذف بيانات مرجعية تاريخية إذا كان خيار التعطيل متاحًا." if rtl else "Do not delete historical master data when Disable is available.",
    ], rtl)

def glossary_tasks(doc, rtl=False):
    heading(doc, "قاموس الحالات والمهام السريعة" if rtl else "Status glossary and quick task finder", 1, rtl)
    rows = [
        ("Available", "متاح للبيع أو للحركة المسموحة." if rtl else "Normally available for the allowed workflow."),
        ("Reserved", "محجوز ولا يتاح للبيع العادي." if rtl else "Reserved and not available for ordinary sale."),
        ("Workshop", "موجود في مسار الورشة." if rtl else "In the Workshop workflow."),
        ("Sold", "تم بيعه وتسجيل أثره المالي." if rtl else "Sold with its financial effect recorded."),
        ("Pending / Approved / In Transit / Received", "مراحل التحويل." if rtl else "Transfer lifecycle states."),
        ("Missing / Unexpected / Variance", "نتائج الجرد." if rtl else "Inventory Count results."),
        ("Draft / Validated / Posted", "مراحل شراء الذهب من العميل." if rtl else "Customer Gold Purchase lifecycle."),
    ]
    table(doc, ["الحالة" if rtl else "Status", "المعنى" if rtl else "Meaning"], rows, rtl, [2.0,4.5])
    heading(doc, "وصفات سريعة" if rtl else "Quick recipes", 2, rtl)
    tasks = [
        "استلام مخزون: المخزون ← إضافة / استلام مخزون ← اختر النوع ← أدخل البيانات ← راجع ← استلام." if rtl else "Receive stock: Inventory -> Add / Receive Inventory -> choose type -> enter data -> review -> Receive.",
        "العثور على قطعة: المخزون ← ابحث بالباركود أو افتح تفاصيل الأصل." if rtl else "Find a piece: Inventory -> search by Barcode or open Asset Details.",
        "التحويل: المخزون ← تحويلات الفروع ← إنشاء ← موافقة ← إرسال ← استلام." if rtl else "Transfer: Inventory -> Transfers -> Create -> Approve -> Dispatch -> Receive.",
        "الجرد: المخزون ← جرد المخزون ← موقع ← بدء ← مسح/عد ← إكمال/إغلاق." if rtl else "Count: Inventory -> Inventory Count -> location -> Start -> scan/count -> Complete/Close.",
        "إنشاء عميل: العملاء ← إضافة عميل ← احفظ ← استخدمه في POS." if rtl else "Create a customer: Customers -> Add Customer -> save -> use it in POS.",
        "بيع: POS ← عميل ← باركود ← سلة ← إجمالي ← دفع ← Checkout مرة واحدة ← إيصال." if rtl else "Sell: POS -> customer -> Barcode -> cart -> totals -> payment -> Checkout once -> receipt.",
        "مراجعة قيد: الحسابات ← ابحث عن المصدر أو المرجع ← راجع المدين والدائن والحالة." if rtl else "Review a journal: Accounting -> find the source/reference -> review debit, credit, and status.",
    ]
    bullets(doc, tasks, rtl)
    heading(doc, "قائمة التسليم والتدريب" if rtl else "Handover and training checklist", 2, rtl)
    checklist = [
        "أعرف كيف أسجل الدخول وأختار الفرع." if rtl else "I know how to sign in and select the correct branch.",
        "أعرف كيف أبحث عن العميل والمورد والقطعة بالباركود." if rtl else "I know how to search for customers, suppliers, and pieces by Barcode.",
        "أعرف كيف أستلم وأحوّل وأرسل للورشة وأنفذ الجرد." if rtl else "I know how to receive, transfer, send to Workshop, and count inventory.",
        "أعرف كيف أراجع الإجمالي والضريبة ومعاينة القيد قبل البيع." if rtl else "I know how to review totals, VAT, and Journal Preview before a sale.",
        "أعرف متى أحتاج صلاحية المسؤول." if rtl else "I know when I need administrator permission.",
        "لن أكرر عملية مالية غير واضحة دون تحقق." if rtl else "I will not repeat an uncertain financial action without verification.",
    ]
    bullets(doc, ["☐ " + x for x in checklist], rtl)
    note(doc, "ورقة تدريب المستخدم السريعة" if rtl else "Quick User Training Checklist", "ابدأ بتسجيل الدخول والفرع، ثم نفذ جولة قراءة في المخزون وPOS والحسابات. لا تستخدم بيانات عميل حقيقية في التدريب الأول." if rtl else "Start with sign-in and branch context, then take a read-only tour of Inventory, POS, and Accounting. Do not use real customer data in the first training session.", rtl, "gold")

def build(language, rtl=False):
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc, language, rtl)
    cover(doc, language, rtl)
    toc(doc, rtl)
    persona_section(doc, rtl)
    basics_section(doc, rtl)
    inventory_section(doc, language, rtl)
    profile_section(doc, language, rtl)
    asset_section(doc, language, rtl)
    lifecycle_section(doc, language, rtl)
    suppliers_section(doc, language, rtl)
    customers_section(doc, language, rtl)
    pos_section(doc, language, rtl)
    sales_section(doc, language, rtl)
    gold_cgp_section(doc, language, rtl)
    finance_section(doc, language, rtl)
    settings_section(doc, language, rtl)
    safety_section(doc, rtl)
    glossary_tasks(doc, rtl)
    note(doc, "ملاحظات الإصدار الحالي" if rtl else "Current release notes", "تكامل الفوترة الإلكترونية الخارجية ليس ضمن نطاق هذا الإصدار. لا تتجاوز حماية السعر أو الضريبة عند ظهور حالة قديمة أو غير متاحة." if rtl else "External e-invoicing integration is not part of this release scope. Do not bypass price or tax protection when a stale or unavailable state appears.", rtl, "gold")
    out = ROOT.parent.parent / ("DARFUS_ERP_End_User_Manual_AR.docx" if rtl else "DARFUS_ERP_End_User_Manual_EN.docx")
    doc.save(out)
    return out

if __name__ == "__main__":
    print(build("ar", True))
    print(build("en", False))
